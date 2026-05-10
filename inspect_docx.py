from docx import Document
import os
import sys

file_path = 'final_report.docx'
if not os.path.exists(file_path):
    print(f'Error: {file_path} not found.')
    sys.exit(1)

doc = Document(file_path)
print('--- Paragraphs ---')
for i, para in enumerate(doc.paragraphs):
    print(f'[{i}] style={para.style.namerash-engineer WMP2902006B} text={para.textrash-engineer WMP2902006B}')
print('
--- Tables ---')
for i, table in enumerate(doc.tables):
    print(f'[TABLE {i}] rows={len(table.rows)} cols={len(table.columns)}')
